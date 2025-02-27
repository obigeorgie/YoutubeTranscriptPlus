import os
import re
import json
import nltk
from urllib.parse import urlparse, parse_qs
from flask import Flask, render_template, request, jsonify, send_file, session, make_response
import youtube_transcript_api
from youtube_transcript_api._errors import TranscriptsDisabled, NoTranscriptFound
import logging
import io
from wordcloud import WordCloud
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from ai_service import AIService

# Download required NLTK data
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('punkt_tab')

# Configure logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET")

ai_service = AIService()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/get-languages', methods=['POST'])
def get_languages():
    try:
        url = request.form.get('url', '')
        logger.info(f"Received request for available languages with URL: {url}")

        if not url:
            logger.warning("No URL provided")
            return jsonify({'error': 'Please enter a YouTube URL'}), 400

        video_id = extract_video_id(url)
        if not video_id:
            logger.warning(f"Invalid YouTube URL: {url}")
            return jsonify({'error': 'Invalid YouTube URL'}), 400

        logger.info(f"Fetching available languages for video ID: {video_id}")
        transcript_list = youtube_transcript_api.YouTubeTranscriptApi.list_transcripts(video_id)

        # Get all available languages
        languages = []
        try:
            # First, identify the original language
            original_language = None
            try:
                original_transcript = transcript_list.find_manually_created_transcript()
                original_language = {
                    'code': original_transcript.language_code,
                    'name': original_transcript.language,
                    'type': 'original'
                }
            except:
                # If no manual transcript, try to get the auto-generated one
                original_transcript = transcript_list.find_generated_transcript()
                original_language = {
                    'code': original_transcript.language_code,
                    'name': original_transcript.language,
                    'type': 'generated'
                }

            # Get manually created transcripts
            for transcript in transcript_list._manually_created_transcripts.values():
                languages.append({
                    'code': transcript.language_code,
                    'name': transcript.language,
                    'type': 'manual'
                })

            # Get auto-generated transcripts
            for transcript in transcript_list._generated_transcripts.values():
                languages.append({
                    'code': transcript.language_code,
                    'name': transcript.language,
                    'type': 'generated'
                })

            # Store original language in session
            if original_language:
                session['original_language'] = original_language

        except Exception as e:
            logger.error(f"Error processing transcripts: {str(e)}")

        if not languages:
            logger.warning(f"No transcripts available for video ID: {video_id}")
            return jsonify({'error': 'No transcripts available for this video'}), 404

        logger.info(f"Found {len(languages)} available languages")
        return jsonify({
            'languages': languages,
            'original_language': original_language,
            'video_id': video_id
        })

    except TranscriptsDisabled:
        logger.warning(f"Transcripts are disabled for video ID: {video_id}")
        return jsonify({'error': 'Transcripts are disabled for this video'}), 400
    except NoTranscriptFound:
        logger.warning(f"No transcript found for video ID: {video_id}")
        return jsonify({'error': 'No transcript found for this video'}), 400
    except Exception as e:
        logger.error(f"Error getting languages: {str(e)}")
        return jsonify({'error': 'An error occurred while fetching available languages'}), 500

@app.route('/get-transcript', methods=['POST'])
def get_transcript():
    try:
        url = request.form.get('url', '')
        language_code = request.form.get('language', 'en')  # Default to English
        logger.info(f"Received request for transcript with URL: {url} and language: {language_code}")

        if not url:
            logger.warning("No URL provided")
            return jsonify({'error': 'Please enter a YouTube URL'}), 400

        video_id = extract_video_id(url)
        if not video_id:
            logger.warning(f"Invalid YouTube URL: {url}")
            return jsonify({'error': 'Invalid YouTube URL'}), 400

        logger.info(f"Fetching transcript for video ID: {video_id} in language: {language_code}")
        transcript_list = youtube_transcript_api.YouTubeTranscriptApi.list_transcripts(video_id)

        try:
            # Try to get the transcript in the requested language
            transcript = transcript_list.find_transcript([language_code])

            # Check if we need to translate
            original_language = session.get('original_language', {}).get('code')
            if original_language and language_code != original_language:
                logger.info(f"Translating transcript from {original_language} to {language_code}")
                transcript = transcript_list.find_transcript([original_language]).translate(language_code)

            transcript_data = transcript.fetch()

        except Exception as e:
            logger.error(f"Error fetching transcript in {language_code}: {str(e)}")
            # Try to get any available transcript if specified language is not available
            transcript = transcript_list.find_transcript([])
            transcript_data = transcript.fetch()
            logger.info(f"Falling back to available transcript in {transcript.language_code}")

        if not transcript_data:
            logger.warning(f"No transcript found for video ID: {video_id}")
            return jsonify({'error': 'No transcript available for this video'}), 404

        logger.info(f"Successfully retrieved transcript for video ID: {video_id}")
        return jsonify({
            'transcript_data': transcript_data,
            'video_id': video_id,
            'language': transcript.language,
            'language_code': transcript.language_code,
            'is_translation': language_code != transcript.language_code
        })

    except TranscriptsDisabled:
        logger.warning(f"Transcripts are disabled for video ID: {video_id}")
        return jsonify({'error': 'Transcripts are disabled for this video'}), 400
    except NoTranscriptFound:
        logger.warning(f"No transcript found for video ID: {video_id}")
        return jsonify({'error': 'No transcript found for this video'}), 400
    except Exception as e:
        logger.error(f"Error getting transcript: {str(e)}")
        return jsonify({'error': 'An error occurred while fetching the transcript'}), 500

def extract_video_id(url):
    try:
        logger.debug(f"Extracting video ID from URL: {url}")
        parsed_url = urlparse(url)

        if 'youtube.com' in parsed_url.netloc:
            if 'watch' in parsed_url.path:
                query_params = parse_qs(parsed_url.query)
                video_id = query_params.get('v', [None])[0]
            elif 'embed' in parsed_url.path or 'v' in parsed_url.path:
                video_id = parsed_url.path.split('/')[-1]
            else:
                video_id = None
        elif 'youtu.be' in parsed_url.netloc:
            video_id = parsed_url.path.lstrip('/')
        else:
            video_id = None

        logger.debug(f"Extracted video ID: {video_id}")
        return video_id
    except Exception as e:
        logger.error(f"Error extracting video ID: {str(e)}")
        return None

@app.route('/download-transcript', methods=['POST'])
def download_transcript():
    try:
        transcript_data = request.form.get('transcript_data', '')  # Get raw transcript data
        format_type = request.form.get('format', 'txt')  # Default to txt if not specified

        if not transcript_data:
            logger.warning("No transcript provided for download")
            return jsonify({'error': 'No transcript to download'}), 400

        try:
            # Parse the transcript data from JSON string
            transcript_entries = json.loads(transcript_data)
        except json.JSONDecodeError:
            # If not JSON, treat as plain text
            transcript_entries = None

        if transcript_entries and isinstance(transcript_entries, list):
            # Generate formatted content based on the requested format
            if format_type == 'srt':
                content = generate_srt(transcript_entries)
                filename = 'transcript.srt'
            elif format_type == 'vtt':
                content = generate_vtt(transcript_entries)
                filename = 'transcript.vtt'
            else:  # Default to plain text
                content = generate_txt(transcript_entries)
                filename = 'transcript.txt'
        else:
            # Fallback to plain text if transcript_data is not in the expected format
            content = transcript_data
            filename = 'transcript.txt'

        # Convert to bytes for sending
        bytes_io = io.BytesIO()
        bytes_io.write(content.encode('utf-8'))
        bytes_io.seek(0)

        logger.info(f"Sending transcript file for download in {format_type} format")
        return send_file(
            bytes_io,
            mimetype='text/plain',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        logger.error(f"Error downloading transcript: {str(e)}")
        return jsonify({'error': 'Failed to download transcript'}), 500

def generate_srt(transcript_entries):
    """Generate SRT format from transcript entries."""
    srt_content = []
    for i, entry in enumerate(transcript_entries, 1):
        start_time = format_timestamp_srt(entry['start'])
        # Calculate end time from start time of next entry or duration
        if i < len(transcript_entries):
            end_time = format_timestamp_srt(transcript_entries[i]['start'])
        else:
            end_time = format_timestamp_srt(entry['start'] + entry.get('duration', 3))

        # Include speaker if available
        text = entry['text'].strip()
        if 'speaker' in entry:
            text = f"{entry['speaker']}: {text}"

        srt_content.extend([
            str(i),
            f"{start_time} --> {end_time}",
            text,
            ""  # Empty line between entries
        ])
    return '\n'.join(srt_content)

def generate_vtt(transcript_entries):
    """Generate VTT format from transcript entries."""
    vtt_content = ["WEBVTT", ""]  # VTT header
    for i, entry in enumerate(transcript_entries):
        start_time = format_timestamp_vtt(entry['start'])
        # Calculate end time from start time of next entry or duration
        if i < len(transcript_entries) - 1:
            end_time = format_timestamp_vtt(transcript_entries[i + 1]['start'])
        else:
            end_time = format_timestamp_vtt(entry['start'] + entry.get('duration', 3))

        # Include speaker if available
        text = entry['text'].strip()
        if 'speaker' in entry:
            text = f"{entry['speaker']}: {text}"

        vtt_content.extend([
            f"{start_time} --> {end_time}",
            text,
            ""  # Empty line between entries
        ])
    return '\n'.join(vtt_content)

def generate_txt(transcript_entries):
    """Generate plain text format from transcript entries."""
    txt_content = []
    for entry in transcript_entries:
        minutes = int(entry['start'] // 60)
        seconds = int(entry['start'] % 60)
        txt_content.append(f"[{minutes:02d}:{seconds:02d}] {entry['text'].strip()}")
    return '\n'.join(txt_content)

def format_timestamp_srt(seconds):
    """Format timestamp for SRT format (HH:MM:SS,mmm)."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds * 1000) % 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

def format_timestamp_vtt(seconds):
    """Format timestamp for VTT format (HH:MM:SS.mmm)."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds * 1000) % 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"

@app.route('/get-word-at-position', methods=['POST'])
def get_word_at_position():
    try:
        x = float(request.form.get('x', 0))
        y = float(request.form.get('y', 0))
        width = float(request.form.get('width', 0))
        height = float(request.form.get('height', 0))

        # Scale coordinates to match the original image size
        x = int((x / width) * 800)  # 800 is the original width
        y = int((y / height) * 400)  # 400 is the original height

        # Get the word at the clicked position from the session
        word_positions = session.get('word_positions', {})
        clicked_word = None

        # Add some tolerance for click detection
        tolerance = 5  # pixels
        for word, positions in word_positions.items():
            if (positions['x'] - tolerance <= x <= positions['x'] + positions['width'] + tolerance and 
                positions['y'] - tolerance <= y <= positions['y'] + positions['height'] + tolerance):
                clicked_word = word
                break

        logger.debug(f"Click at coordinates: ({x}, {y}), Found word: {clicked_word}")
        return jsonify({'word': clicked_word})
    except Exception as e:
        logger.error(f"Error getting word at position: {str(e)}")
        return jsonify({'error': 'Failed to get word at position'}), 500

def generate_wordcloud():
    try:
        transcript_data = request.form.get('transcript_data', '')
        if not transcript_data:
            return jsonify({'error': 'No transcript data provided'}), 400

        # Parse transcript data
        try:
            transcript_entries = json.loads(transcript_data)
        except json.JSONDecodeError:
            return jsonify({'error': 'Invalid transcript data format'}), 400

        # Combine all text from transcript
        full_text = ' '.join(entry['text'] for entry in transcript_entries)

        # Basic text cleaning
        full_text = re.sub(r'[^\w\s]', '', full_text)
        full_text = ' '.join(full_text.split())

        # Filter out stop words
        stop_words = set(stopwords.words('english'))
        words = [word.lower() for word in full_text.split() 
                if word.lower() not in stop_words and word.isalpha()]

        if not words:
            return jsonify({'error': 'No valid words found for word cloud'}), 400

        # Generate word cloud with position tracking
        wordcloud = WordCloud(
            width=800,
            height=400,
            background_color='#2d2d2d',
            colormap='viridis',
            max_words=100,
            min_font_size=10,
            max_font_size=60
        )

        # Generate the word cloud
        wordcloud.generate(' '.join(words))

        # Get word positions with improved accuracy
        word_positions = {}
        for (word, freq), font_size, position, orientation, color in wordcloud.layout_:
            x, y = position
            width = len(word) * (font_size / 2)  # More accurate width calculation
            height = font_size * 1.2  # Add some padding
            word_positions[word] = {
                'x': int(x),
                'y': int(y - height/2),  # Center the clickable area
                'width': int(width),
                'height': int(height)
            }
            logger.debug(f"Word '{word}' position: x={x}, y={y}, width={width}, height={height}")

        # Store positions in session for click handling
        session['word_positions'] = word_positions

        # Convert to image
        img_io = io.BytesIO()
        wordcloud.to_image().save(img_io, 'PNG')
        img_io.seek(0)

        return send_file(
            img_io,
            mimetype='image/png',
            as_attachment=False
        )

    except Exception as e:
        logger.error(f"Error generating word cloud: {str(e)}")
        return jsonify({'error': 'Failed to generate word cloud'}), 500

@app.route('/generate-wordcloud', methods=['POST'])
def generate_wordcloud_route():
    return generate_wordcloud()


@app.route('/analyze-transcript', methods=['POST'])
def analyze_transcript():
    try:
        transcript_data = request.form.get('transcript_data', '')
        analysis_type = request.form.get('type', '')  # 'summary' or 'key_points'

        if not transcript_data:
            return jsonify({'error': 'No transcript data provided'}), 400

        try:
            transcript_entries = json.loads(transcript_data)
        except json.JSONDecodeError:
            return jsonify({'error': 'Invalid transcript data format'}), 400

        # Combine all text from transcript
        full_text = ' '.join(entry['text'] for entry in transcript_entries)

        try:
            if analysis_type == 'summary':
                result = ai_service.summarize_transcript(full_text)
                return jsonify({'summary': result})
            elif analysis_type == 'key_points':
                result = ai_service.extract_key_points(full_text)
                return jsonify({'key_points': result})
            else:
                return jsonify({'error': 'Invalid analysis type'}), 400
        except Exception as e:
            logger.error(f"Error in AI analysis: {e}")
            return jsonify({'error': 'Failed to analyze transcript'}), 500

    except Exception as e:
        logger.error(f"Error in analyze_transcript: {e}")
        return jsonify({'error': 'Failed to process request'}), 500

@app.route('/identify-speakers', methods=['POST'])
def identify_speakers():
    try:
        transcript_data = request.form.get('transcript_data', '')

        if not transcript_data:
            return jsonify({'error': 'No transcript data provided'}), 400

        try:
            transcript_segments = json.loads(transcript_data)
        except json.JSONDecodeError:
            return jsonify({'error': 'Invalid transcript data format'}), 400

        try:
            # Use AI service to identify speakers
            identified_segments = ai_service.identify_speakers(transcript_segments)
            return jsonify({'segments': identified_segments})
        except Exception as e:
            logger.error(f"Error in speaker identification: {e}")
            return jsonify({'error': 'Failed to identify speakers'}), 500

    except Exception as e:
        logger.error(f"Error in identify_speakers route: {e}")
        return jsonify({'error': 'Failed to process request'}), 500

@app.route('/export-transcript', methods=['POST'])
def export_transcript():
    try:
        transcript_data = request.form.get('transcript_data', '')
        format_type = request.form.get('format', 'txt')  # Default to txt if not specified
        video_id = request.form.get('video_id', '')
        title = request.form.get('title', 'Transcript')

        if not transcript_data:
            logger.warning("No transcript provided for export")
            return jsonify({'error': 'No transcript to export'}), 400

        try:
            # Parse the transcript data from JSON string
            transcript_entries = json.loads(transcript_data)
        except json.JSONDecodeError:
            transcript_entries = None

        if transcript_entries and isinstance(transcript_entries, list):
            # Generate formatted content based on the requested format
            if format_type == 'srt':
                content = generate_srt(transcript_entries)
                filename = f'transcript_{video_id}.srt'
                mimetype = 'text/plain'
            elif format_type == 'vtt':
                content = generate_vtt(transcript_entries)
                filename = f'transcript_{video_id}.vtt'
                mimetype = 'text/plain'
            elif format_type == 'txt':
                content = generate_txt(transcript_entries)
                filename = f'transcript_{video_id}.txt'
                mimetype = 'text/plain'
            elif format_type == 'html':
                content = generate_html(transcript_entries, title)
                filename = f'transcript_{video_id}.html'
                mimetype = 'text/html'
            elif format_type == 'pdf':
                pdf_bytes = generate_pdf(transcript_entries, title)
                return send_file(
                    io.BytesIO(pdf_bytes),
                    mimetype='application/pdf',
                    as_attachment=True,
                    download_name=f'transcript_{video_id}.pdf'
                )
            elif format_type == 'docx':
                docx_bytes = generate_docx(transcript_entries, title)
                return send_file(
                    io.BytesIO(docx_bytes),
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True,
                    download_name=f'transcript_{video_id}.docx'
                )
            else:
                return jsonify({'error': 'Unsupported format'}), 400
        else:
            content = transcript_data
            filename = f'transcript_{video_id}.txt'
            mimetype = 'text/plain'

        # Send the file
        response = make_response(content)
        response.headers['Content-Type'] = mimetype
        response.headers['Content-Disposition'] = f'attachment; filename={filename}'
        return response

    except Exception as e:
        logger.error(f"Error exporting transcript: {str(e)}")
        return jsonify({'error': 'Failed to export transcript'}), 500

def generate_html(transcript_entries, title):
    """Generate HTML format from transcript entries."""
    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{title}</title>
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 2rem; }}
            .transcript-entry {{ margin-bottom: 1rem; }}
            .timestamp {{ color: #666; font-family: monospace; }}
            .speaker {{ font-weight: bold; color: #2196F3; }}
        </style>
    </head>
    <body>
        <h1>{title}</h1>
        <div class="transcript">
    """

    for entry in transcript_entries:
        minutes = int(entry['start'] // 60)
        seconds = int(entry['start'] % 60)
        timestamp = f"{minutes:02d}:{seconds:02d}"

        speaker_html = f'<span class="speaker">{entry["speaker_id"]}: </span>' if 'speaker_id' in entry else ''

        html_content += f"""
            <div class="transcript-entry">
                <span class="timestamp">[{timestamp}]</span>
                {speaker_html}
                <span class="text">{entry['text']}</span>
            </div>
        """

    html_content += """
        </div>
    </body>
    </html>
    """
    return html_content

def generate_pdf(transcript_entries, title):
    """Generate PDF format from transcript entries."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from io import BytesIO

        # Create PDF buffer
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()

        # Create custom styles
        styles.add(ParagraphStyle(
            name='TranscriptEntry',
            parent=styles['Normal'],
            spaceAfter=10,
            leftIndent=20
        ))

        # Build PDF content
        content = []
        content.append(Paragraph(title, styles['Title']))
        content.append(Spacer(1, 12))

        for entry in transcript_entries:
            minutes = int(entry['start'] // 60)
            seconds = int(entry['start'] % 60)
            timestamp = f"{minutes:02d}:{seconds:02d}"

            speaker_text = f"{entry['speaker_id']}: " if 'speaker_id' in entry else ''
            text = f"[{timestamp}] {speaker_text}{entry['text']}"

            content.append(Paragraph(text, styles['TranscriptEntry']))

        # Build PDF
        doc.build(content)
        return buffer.getvalue()

    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        raise

def generate_docx(transcript_entries, title):
    """Generate DOCX format from transcript entries."""
    try:
        from docx import Document
        from docx.shared import Pt
        from io import BytesIO

        # Create document
        doc = Document()
        doc.add_heading(title, 0)

        # Add transcript entries
        for entry in transcript_entries:
            minutes = int(entry['start'] // 60)
            seconds = int(entry['start'] % 60)
            timestamp = f"{minutes:02d}:{seconds:02d}"

            speaker_text = f"{entry['speaker_id']}: " if 'speaker_id' in entry else ''
            text = f"[{timestamp}] {speaker_text}{entry['text']}"

            paragraph = doc.add_paragraph(text)
            paragraph.style.font.size = Pt(11)

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        raise

@app.route('/generate-share-link', methods=['POST'])
def generate_share_link():
    try:
        video_id = request.form.get('video_id', '')
        timestamp = request.form.get('timestamp', '')

        if not video_id:
            return jsonify({'error': 'No video ID provided'}), 400

        # Generate a shareable YouTube link
        share_link = f'https://youtu.be/{video_id}'
        if timestamp:
            share_link += f'?t={int(float(timestamp))}'

        return jsonify({'share_link': share_link})

    except Exception as e:
        logger.error(f"Error generating share link: {str(e)}")
        return jsonify({'error': 'Failed to generate share link'}), 500

if __name__ == '__main__':
    app.run(debug=True)